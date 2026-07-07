import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ─── 样式设置 ───
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

base_dir = r"D:\java\SmartLightingExp\docs"
prototype_dir = os.path.join(base_dir, "移动端原型图")

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return h

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table

def add_image(name, width_inches=5.5):
    path = os.path.join(prototype_dir, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[图片缺失: {name}]")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

def add_bold(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    return p

# ===========================================
# 封面
# ===========================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("智慧路灯管理系统\n移动端接口文档")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("\n第一阶段：数字孪生 / 数据报表 / 设备管理 / 智能助手")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

# ===========================================
# 一、通用说明
# ===========================================
add_heading("一、通用说明", 1)

add_heading("1.1 接口地址", 2)
doc.add_paragraph("当前阶段通过 natapp 隧道访问。地址见《移动端开发参考.md》，本文档中所有接口路径使用占位符表示：")
add_code("{{BASE_URL}}")
doc.add_paragraph("移动端开发时，将 {{BASE_URL}} 替换为实际地址即可（如 http://47.96.27.141:8080），只需改一处全局常量。")

add_heading("1.2 统一响应格式", 2)
add_code('{ "code": 200, "msg": "success", "data": { ... } }')
add_table(
    ["code", "含义"],
    [["200", "成功"], ["400", "参数错误"], ["401", "未登录 / Token过期"],
     ["403", "无权限"], ["404", "资源不存在"], ["409", "数据冲突"], ["500", "服务器错误"]]
)

add_heading("1.3 认证方式", 2)
doc.add_paragraph("除登录接口外，所有业务接口需在请求头携带 JWT Token：")
add_code("Authorization: Bearer <token>")
doc.add_paragraph("Token 通过登录接口获取，有效期 24 小时，过期后需重新登录。")

add_heading("1.4 分页规范", 2)
doc.add_paragraph("POST 分页请求体：")
add_code('{ "page": 1, "size": 20, "...查询条件": "" }')
doc.add_paragraph("GET 分页 URL：")
add_code("?pageNum=1&pageSize=20&keyword=xxx")
doc.add_paragraph("分页响应：")
add_code('{ "records": [...], "total": 100, "size": 20, "current": 1, "pages": 5 }')

# ===========================================
# 二、认证模块
# ===========================================
add_heading("二、认证模块", 1)
doc.add_paragraph("对应原型图：loginscreen.png")
add_image("loginscreen.png", 3.0)

add_heading("2.1 登录", 2)
add_code("POST {{BASE_URL}}/api/auth/login")
add_code('{ "username": "admin", "password": "admin123" }')
doc.add_paragraph("响应 data 字段：")
add_table(
    ["字段", "类型", "说明"],
    [["token", "String", "JWT Token"], ["username", "String", "用户名"],
     ["roleCode", "String", "角色编码"], ["permissions", "String[]", "权限码列表"],
     ["menus", "Array", "可见菜单树"]]
)
doc.add_paragraph("测试账号：admin / admin123（SUPER_ADMIN）")

add_heading("2.2 校验 Token", 2)
add_code("GET {{BASE_URL}}/api/auth/me")
doc.add_paragraph("App 启动时调用，校验 Token 是否有效。返回结构与登录接口相同。")

# ===========================================
# 三、设备管理
# ===========================================
add_heading("三、设备管理模块", 1)
doc.add_paragraph("对应原型图：devicedetailscreen.png")
add_image("devicedetailscreen.png")

add_heading("3.1 设备列表（POST，推荐）", 2)
add_code("POST {{BASE_URL}}/api/devices/list")
add_table(
    ["参数", "类型", "必填", "说明"],
    [["page", "int", "否", "页码，默认1"], ["size", "int", "否", "每页条数，默认20"],
     ["deviceId", "String", "否", "设备编号（精确）"], ["name", "String", "否", "名称（模糊）"],
     ["area", "String", "否", "区域名称"], ["areaId", "Long", "否", "区域ID"],
     ["status", "Integer", "否", "0停用/1在线/2离线/3异常"],
     ["enabled", "Boolean", "否", "是否启用"], ["healthScoreMin", "BigDecimal", "否", "健康分下限"],
     ["healthScoreMax", "BigDecimal", "否", "健康分上限"]]
)
doc.add_paragraph("响应 records 主要字段：deviceId, name, area, areaId, location, status, healthScore, lastHeartbeatAt, latestData, firmwareVersion, enabled")

add_heading("3.2 设备列表（GET）", 2)
add_code("GET {{BASE_URL}}/api/devices/page?pageNum=1&pageSize=20&keyword=xxx&area=A区&status=1")

add_heading("3.3 设备详情", 2)
add_code("GET {{BASE_URL}}/api/devices/{deviceId}")

add_heading("3.4 新增设备", 2)
add_code("POST {{BASE_URL}}/api/devices")
add_code('{"deviceId":"SL-008","name":"学院路8号灯杆","area":"学院路","areaId":3,"location":"29.5647,106.5478","status":1,"enabled":true}')

add_heading("3.5 编辑设备", 2)
add_code("PUT {{BASE_URL}}/api/devices/{deviceId}")

add_heading("3.6 删除设备（软删除）", 2)
add_code("DELETE {{BASE_URL}}/api/devices/{deviceId}")

add_heading("3.7 批量新增", 2)
add_code("POST {{BASE_URL}}/api/devices/batch")
add_code('[{"deviceId":"SL-009","name":"设备9","area":"A区","status":1}, ...]')

add_heading("3.8 批量分配区域", 2)
add_code("PUT {{BASE_URL}}/api/devices/batch-area")
add_code('{"deviceIds":[1,2,3],"areaId":5}')

# ===========================================
# 四、设备健康评分
# ===========================================
add_heading("四、设备健康评分", 1)

add_heading("4.1 健康评分详情（实时计算）", 2)
add_code("GET {{BASE_URL}}/api/devices/{deviceId}/health")
doc.add_paragraph("响应：overallScore(0-100)、level(优秀/良好/一般/较差/危险)、levelColor、dimensions(四个维度: 离线频次30%/通信质量25%/指令响应率25%/传感器状态20%)、suggestion")

add_heading("4.2 全局健康概览", 2)
add_code("GET {{BASE_URL}}/api/devices/health/summary")
doc.add_paragraph("响应：totalDevices、healthyCount、warningCount、criticalCount、averageScore、list")

# ===========================================
# 五、遥测数据
# ===========================================
add_heading("五、设备遥测数据", 1)

add_heading("5.1 最新遥测", 2)
add_code("GET {{BASE_URL}}/api/telemetry/latest/{deviceId}")
doc.add_paragraph("响应 data 对象：illuminance(lux)、temperature(°C)、humidity(%)、pm25、aqi、pir(0无人/1有人)、traffic")

add_heading("5.2 遥测历史（分页）", 2)
add_code("POST {{BASE_URL}}/api/telemetry/history")
add_code('{"deviceId":"SL-001","collectedAtFrom":"2026-07-06T00:00:00","collectedAtTo":"2026-07-06T23:59:59","page":1,"size":50}')

# ===========================================
# 六、设备控制
# ===========================================
add_heading("六、设备控制模块", 1)
doc.add_paragraph("对应原型图：manualcontrolscreen.png、controllogscreen.png")
add_image("manualcontrolscreen.png")
add_image("controllogscreen.png")

add_heading("6.1 下发控制指令", 2)
add_code("POST {{BASE_URL}}/api/devices/{deviceId}/control")
add_table(
    ["action", "说明", "brightness 是否必填"],
    [["ON", "开灯", "否"], ["OFF", "关灯", "否"], ["DIMMING", "调光", "是（0-100）"]]
)
doc.add_paragraph("手动锁定机制：手动控制后 AI 策略引擎 30 分钟内跳过该设备，避免自动覆盖。")

add_heading("6.2 解除手动锁定", 2)
add_code("DELETE {{BASE_URL}}/api/devices/{deviceId}/manual-lock")

add_heading("6.3 控制历史", 2)
add_code("GET {{BASE_URL}}/api/devices/{deviceId}/control-history?page=1&size=10")
doc.add_paragraph("响应字段：action(ON/OFF/DIMMING(80))、brightness、source(MANUAL/AUTO)、operator、status(SENT/SUCCESS/FAILED)、issuedAt、ackAt、resultDetail")

# ===========================================
# 七、数字孪生
# ===========================================
add_heading("七、数字孪生（仪表盘）模块", 1)
doc.add_paragraph("对应原型图：mapscreen.png")
add_image("mapscreen.png")

add_heading("7.1 统计概览", 2)
add_code("GET {{BASE_URL}}/api/dashboard/stats")
doc.add_paragraph("响应：totalDevices、onlineDevices、onlineRate、alertCount、energySavingRate、todayEnergy")

add_heading("7.2 能耗趋势", 2)
add_code("GET {{BASE_URL}}/api/dashboard/energy-trend")
doc.add_paragraph("响应：labels(24小时)、current(今日各时段kWh)、lastWeek(上周同期)")

add_heading("7.3 分区设备状态", 2)
add_code("GET {{BASE_URL}}/api/dashboard/districts")
doc.add_paragraph("响应数组：name、online、offline、warning、disabled")

add_heading("7.4 边缘AI决策状态", 2)
add_code("GET {{BASE_URL}}/api/dashboard/edge-status")
doc.add_paragraph("响应：totalDecisions、hitCount、lastSimulatedAt、enabled")

add_heading("7.5 边缘AI最近决策", 2)
add_code("GET {{BASE_URL}}/api/dashboard/edge/recent")

add_heading("7.6 手动触发边缘模拟", 2)
add_code("POST {{BASE_URL}}/api/dashboard/edge/trigger")

# ===========================================
# 八、数据报表
# ===========================================
add_heading("八、数据报表模块", 1)
doc.add_paragraph("对应原型图：energyanalysisscreen.png")
add_image("energyanalysisscreen.png")

add_heading("8.1 年度能耗统计", 2)
add_code("GET {{BASE_URL}}/api/dashboard/energy/yearly-stats?year=2026")
doc.add_paragraph("响应：year、totalKwh、savedKwh、carbonReductionKg(kg CO2)、avgOnlineRate、lastYear")

add_heading("8.2 月度能耗（12个月）", 2)
add_code("GET {{BASE_URL}}/api/dashboard/energy/monthly?year=2026")
doc.add_paragraph("响应：months([1月...12月])、consumption[]、savings[]")

add_heading("8.3 分区分区能耗占比", 2)
add_code("GET {{BASE_URL}}/api/dashboard/energy/district?year=2026")
doc.add_paragraph("响应数组：name(区域)、value(kWh)")

# ===========================================
# 九、智能助手
# ===========================================
add_heading("九、智能助手模块", 1)
doc.add_paragraph("对应原型图：smartassistantscreen.png")
add_image("smartassistantscreen.png")

add_heading("9.1 对话问答", 2)
add_code("POST {{BASE_URL}}/api/assistant/chat")
add_code('{"message":"灯不亮怎么办"}')
doc.add_paragraph("响应 type：KNOWLEDGE_QA(知识问答) / THRESHOLD_UPDATED(参数已修改)")
doc.add_paragraph("自然语言调参示例：'把阈值调到30'、'亮度调到60%'")

add_heading("9.2 一键诊断", 2)
add_code("POST {{BASE_URL}}/api/assistant/diagnose")
add_code('{"deviceId":"SL-001","question":"最近频繁离线是什么原因"}')
doc.add_paragraph("AI 结合最新遥测/告警/控制历史做综合分析，返回同聊天结构。")

# ===========================================
# 十、告警中心（配套）
# ===========================================
add_heading("十、告警中心模块（配套）", 1)
doc.add_paragraph("对应原型图：alarmcenterscreen.png")
add_image("alarmcenterscreen.png")

add_heading("10.1 告警列表", 2)
add_code("POST {{BASE_URL}}/api/alarms/list")
add_code('{"deviceId":"SL-001","type":"OFFLINE","level":"HIGH","status":"ACTIVE","page":1,"size":20}')
doc.add_paragraph("告警类型(OFFLINE/HIGH_TEMP/LOW_HEALTH)、级别(HIGH/MEDIUM/LOW)、状态(ACTIVE/ACKNOWLEDGED/RESOLVED)")

add_heading("10.2 告警详情", 2)
add_code("GET {{BASE_URL}}/api/alarms/{id}")

add_heading("10.3 处理确认告警", 2)
add_code("PUT {{BASE_URL}}/api/alarms/{id}/handle")
add_code('{"remark":"设备已恢复正常"}')

add_heading("10.4 告警统计", 2)
add_code("GET {{BASE_URL}}/api/alarms/stats")

add_heading("10.5 告警趋势", 2)
add_code("GET {{BASE_URL}}/api/alarms/trend?days=7")

# ===========================================
# 十一、对照表
# ===========================================
add_heading("十一、原型图与接口对照表", 1)
add_table(
    ["原型图", "对应模块", "涉及接口"],
    [["loginscreen.png", "登录认证", "2.1"],
     ["devicedetailscreen.png", "设备详情", "3.3, 4.1, 5.1"],
     ["manualcontrolscreen.png", "手动控制", "6.1, 6.2"],
     ["controllogscreen.png", "控制历史", "6.3"],
     ["energyanalysisscreen.png", "数据报表", "8.1, 8.2, 8.3"],
     ["smartassistantscreen.png", "智能助手", "9.1, 9.2"],
     ["mapscreen.png", "数字孪生", "7.1~7.6, 3.1"],
     ["alarmcenterscreen.png", "告警中心", "10.1~10.5"],
     ["strategysettingsscreen.png", "策略设置", "不在第一阶段"]]
)

add_heading("十二、接口覆盖评估", 1)
add_table(
    ["模块", "所需接口数", "已就绪", "缺失"],
    [["登录认证", "2", "2", "0"], ["设备管理(含健康/遥测/控制)", "15", "15", "0"],
     ["数字孪生(仪表盘)", "6", "6", "0"], ["数据报表", "3", "3", "0"],
     ["智能助手", "2", "2", "0"], ["告警中心(配套)", "5", "5", "0"],
     ["合计", "33", "33", "0"]]
)
doc.add_paragraph("✅ 当前后端 API 完整覆盖移动端第一阶段所有需求，无缺失接口。")

# ===========================================
# 十三、BASE_URL 占位符说明
# ===========================================
add_heading("十三、关于 {{BASE_URL}} 占位符", 1)
doc.add_paragraph("本文档及《移动端开发参考.md》中所有接口地址均使用 {{BASE_URL}} 占位符，而非硬编码的 natapp 域名。")
doc.add_paragraph("原因：natapp 免费隧道每次重启域名可能变化，使用占位符便于移动端只需改动一处全局常量。")
doc.add_paragraph("移动端使用方式：")

add_code("// Flutter")
add_code('const BASE_URL = "http://实际地址:端口";')
add_code("// Android/Java")
add_code('public static final String BASE_URL = "http://实际地址:端口";')
add_code("// iOS/Swift")
add_code('let BASE_URL = "http://实际地址:端口"')
add_code("// React Native")
add_code('const BASE_URL = "http://实际地址:端口";')

doc.add_paragraph("当前联调阶段，实际地址由后端同学（natapp 隧道）提供。")
doc.add_paragraph("后端部署到云服务器后，改为固定 IP：http://47.96.27.141:8080")

# ===========================================
# 保存
# ===========================================
output_path = os.path.join(base_dir, "移动端接口文档.docx")
doc.save(output_path)
print(f"Word document generated: {output_path}")
